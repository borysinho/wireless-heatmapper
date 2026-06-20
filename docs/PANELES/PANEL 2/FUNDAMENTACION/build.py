#!/usr/bin/env python3
from __future__ import annotations

import argparse
import re
import shutil
import subprocess
import sys
from pathlib import Path

import lxml.etree as ET
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

INCLUDE_RE = re.compile(r"\[\[include:\s*(.*?)\s*\]\]")
PLANTUML_RE = re.compile(r"```plantuml\n(.*?)\n```", re.DOTALL)
TITLE_RE = re.compile(r"^\s*title\s+(.+?)\s*$", re.MULTILINE)
STARTUML_ID_RE = re.compile(r"@startuml[ \t]+(\S+)")
HEADING_RE = re.compile(r"^(#{1,5}) (.+)$", re.MULTILINE)
SECTION_NUM_RE = re.compile(r"^(\d+)((?:\.\d+)*?)\.?\s+(.*)")

# Mapa de números del Panel 1 → números del Capítulo 1 del Panel 2
PANEL1_SECTION_MAP: dict[str, str] = {
    "3": "1.1", "4": "1.2", "5": "1.3", "6": "1.4", "7": "1.5",
    "8": "1.6", "9": "1.7", "10": "1.8", "11": "1.9", "12": "1.10",
}
PAGE_BREAK = """```{=openxml}
<w:p><w:r><w:br w:type="page"/></w:r></w:p>
```"""

PANEL_DIR = Path(__file__).resolve().parent
INDEX_FILE = PANEL_DIR / "00-indice.md"
REFERENCE_DOC = PANEL_DIR / "reference.docx"
IMG_DIR = PANEL_DIR / "img"
REPO_ROOT = PANEL_DIR.parents[3]
WORKSPACE = REPO_ROOT / ".panel_build" / "panel2"
OUTPUT_FILE = PANEL_DIR.parent / "Wireless-HeatMapper-Panel2.docx"

DIAGRAMS_PNG_DIR = REPO_ROOT / "diagrams" / "paneles" / "png"

# Mapeo: identificador del bloque PlantUML → prefijo del PNG exportado desde StarUML.
# Los diagramas sin entrada en este mapa (o cuyo PNG no exista) quedan en blanco.
DIAGRAM_PNG_MAP: dict[str, str] = {
    # Panel 1 — descripción del problema
    "__mindmap__":                    "p1-01-ishikawa-causa-efecto",
    "Modelo_Dominio_WirelessHeatMapper": "p1-02-modelo-dominio",
    "Situacion_Actual_Deseada":       "p1-12-situacion-actual-deseada",
    # Panel 1 — Sprint 0
    "Modelo_Contexto_Sprint0":        "p1-s0-01-casos-uso",
    "Arquitectura_Paquetes_Sprint0":  "p1-s0-02-paquetes",
    "Despliegue_Sprint0":             "p1-s0-03-despliegue",
    "Modelo_Datos_Conceptual_Sprint0": "p1-s0-04-datos-conceptual",
    # Panel 1 — Sprint 1
    "Contexto_Sprint1":               "p1-s1-01-casos-uso",
    "Paquetes_Sprint1":               "p1-s1-02-paquetes",
    "Despliegue_Sprint1":             "p1-s1-03-despliegue",
    "Datos_Conceptual_Sprint1":       "p1-s1-04-datos-conceptual",
    "Secuencia_Autenticacion_Sprint1": "p1-s1-05-secuencia-autenticacion",
    # Plan de sprints
    "__pert__":                       "pert-dependencias",
    # Panel 2 — Sprint 2
    "Sprint2_CasosUso":               "p2-s2-01-casos-uso",
    "Sprint2_Clases":                 "p2-s2-02-clases",
    "Sprint2_Secuencia_Subida":       "p2-s2-03-secuencia-subida",
    "Sprint2_Secuencia_Calibracion":  "p2-s2-04-secuencia-calibracion",
    # Panel 2 — Sprint 3
    "Sprint3_CasosUso":               "p2-s3-01-casos-uso",
    "Sprint3_Secuencia_Captura":      "p2-s3-02-secuencia-captura",
    "Sprint3_Estados":                "p2-s3-03-estados-captura",
}


def _add_compact_style(doc: Document) -> None:
    """Add pandoc's Compact paragraph style for table cells (required for LibreOffice compatibility)."""
    existing_ids = {
        s.element.get(qn("w:styleId"))
        for s in doc.styles
        if s.element.get(qn("w:styleId"))
    }
    if "Compact" in existing_ids:
        return
    styles_elem = doc.part.styles.element
    compact_xml = (
        '<w:style xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
        'w:customStyle="1" w:styleId="Compact" w:type="paragraph">'
        '<w:name w:val="Compact"/>'
        '<w:basedOn w:val="BodyText"/>'
        "<w:qFormat/>"
        "<w:pPr><w:spacing w:after=\"36\" w:before=\"36\"/></w:pPr>"
        "</w:style>"
    )
    styles_elem.append(ET.fromstring(compact_xml))


def ensure_reference_doc() -> None:
    REFERENCE_DOC.parent.mkdir(parents=True, exist_ok=True)
    IMG_DIR.mkdir(parents=True, exist_ok=True)

    if REFERENCE_DOC.exists():
        # Patch existing reference.docx if it lacks the Compact style.
        doc = Document(REFERENCE_DOC)
        _add_compact_style(doc)
        doc.save(REFERENCE_DOC)
        return

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)

    for style_name, size in (("Heading 1", 14), ("Heading 2", 13), ("Heading 3", 12)):
        style = doc.styles[style_name]
        style.font.name = "Times New Roman"
        style.font.size = Pt(size)
        style.font.bold = True

    if "Table Grid Copilot" not in [style.name for style in doc.styles]:
        table_style = doc.styles.add_style("Table Grid Copilot", WD_STYLE_TYPE.TABLE)
        table_style.base_style = doc.styles["Table Grid"]

    _add_compact_style(doc)
    doc.save(REFERENCE_DOC)


def parse_includes() -> list[Path]:
    text = INDEX_FILE.read_text(encoding="utf-8")
    includes: list[Path] = []
    for match in INCLUDE_RE.finditer(text):
        include_path = (INDEX_FILE.parent / match.group(1)).resolve()
        if not include_path.exists():
            raise FileNotFoundError(f"No existe el archivo incluido: {include_path}")
        includes.append(include_path)
    return includes


def _resolve_diagram_key(block: str) -> str | None:
    """Devuelve la clave de búsqueda en DIAGRAM_PNG_MAP para un bloque PlantUML.

    Retorna None para bloques Gantt (siempre en blanco) y para bloques
    @startuml sin identificador conocido.
    """
    if re.search(r"@startgantt", block):
        return None
    if re.search(r"@startmindmap", block):
        return "__mindmap__"
    id_match = STARTUML_ID_RE.search(block)
    if id_match:
        return id_match.group(1)
    # @startuml sin ID: usar el título para distinguir casos conocidos
    title_match = TITLE_RE.search(block)
    if title_match and "PERT" in title_match.group(1):
        return "__pert__"
    return None


def _find_png_for_stem(stem: str) -> Path | None:
    """Busca el PNG exportado de mayor índice para el prefijo dado."""
    candidates = list(DIAGRAMS_PNG_DIR.glob(f"{stem}!*.png"))
    if not candidates:
        return None

    def _export_number(p: Path) -> int:
        m = re.search(r"_(\d+)\.png$", p.name)
        return int(m.group(1)) if m else -1

    return max(candidates, key=_export_number)


def render_plantuml(markdown: str, image_dir: Path, counter: list[int]) -> str:
    """Reemplaza bloques ```plantuml con el PNG exportado desde StarUML.

    Si el diagrama no tiene PNG asignado o el archivo no existe, el bloque
    se reemplaza con una línea en blanco (el espacio queda vacío en el docx).
    """
    def replace(match: re.Match[str]) -> str:
        counter[0] += 1
        block = match.group(1)
        title_match = TITLE_RE.search(block)
        title = title_match.group(1).strip() if title_match else f"Diagrama {counter[0]}"

        key = _resolve_diagram_key(block)
        if key is None or key not in DIAGRAM_PNG_MAP:
            return ""

        stem = DIAGRAM_PNG_MAP[key]
        source_png = _find_png_for_stem(stem)
        if source_png is None:
            return ""

        dest_png = image_dir / source_png.name
        shutil.copy2(source_png, dest_png)
        return f"![{title}](img/{dest_png.name})"

    return PLANTUML_RE.sub(replace, markdown)


def strip_trailing_separator(text: str) -> str:
    return re.sub(r"\n---\s*$", "", text.strip())


def is_panel1_content(path: Path) -> bool:
    """True para archivos narrativos del Panel 1 que deben ser demotados un nivel en el Panel 2."""
    p = str(path)
    return (
        "PANEL 1" in p
        and "PERFIL-PROYECTO" in p
        and "SPRINT-" not in p
        and "01-caratula" not in p
    )


def demote_headings(content: str) -> str:
    """Demota todos los headings un nivel y remapea la numeración del Panel 1 al Capítulo 1.

    Ejemplos:
        # 3. Introducción           →  ## 1.1 Introducción
        ## 4.1 Fundamentación       →  ### 1.2.1 Fundamentación Teórica
        ### 4.1.1 Propagación       →  #### 1.2.1.1 Propagación de señal RF
        ### Resumen comparativo     →  #### Resumen comparativo  (sin número, solo demota)
    """
    def _remap_title(title: str) -> str:
        m = SECTION_NUM_RE.match(title)
        if not m:
            return title  # sin prefijo numérico — solo demotar
        major, sub, text = m.group(1), m.group(2), m.group(3)
        if major not in PANEL1_SECTION_MAP:
            return title
        return f"{PANEL1_SECTION_MAP[major]}{sub} {text}"

    def _replace(m: re.Match) -> str:
        hashes, title = m.group(1), m.group(2)
        return f"{hashes}# {_remap_title(title)}"

    return HEADING_RE.sub(_replace, content)


def assemble_markdown(includes: list[Path], workspace: Path) -> tuple[Path, int]:
    workspace.mkdir(parents=True, exist_ok=True)
    img_workspace = workspace / "img"
    img_workspace.mkdir(parents=True, exist_ok=True)

    # Copiar imágenes Gantt exportadas desde StarUML
    gantt_images_dir = REPO_ROOT / "diagrams" / "paneles" / "gantt" / "images"
    if gantt_images_dir.exists():
        for img in sorted(gantt_images_dir.glob("*.png")):
            shutil.copy2(img, img_workspace / img.name)

    # Copiar capturas de software similar (antecedentes sección 4.2, incluida desde Panel 1)
    similar_sw_dir = REPO_ROOT / "docs" / "PANELES" / "PANEL 1" / "PERFIL-PROYECTO" / "img" / "similar-sw"
    if similar_sw_dir.exists():
        for img in sorted(similar_sw_dir.glob("*.png")):
            shutil.copy2(img, img_workspace / img.name)

    counter = [0]
    chunks: list[str] = []
    for index, include in enumerate(includes):
        content = include.read_text(encoding="utf-8")
        content = render_plantuml(content, img_workspace, counter)
        if is_panel1_content(include):
            content = demote_headings(content)
        chunks.append(strip_trailing_separator(content))
        if index != len(includes) - 1:
            chunks.append(PAGE_BREAK)

    assembled = workspace / "assembled.md"
    assembled.write_text("\n\n".join(chunks) + "\n", encoding="utf-8")
    shutil.copy2(REFERENCE_DOC, workspace / "reference.docx")
    return assembled, counter[0]


def run_pandoc(assembled: Path, workspace: Path, output_file: Path) -> None:
    command = [
        "pandoc",
        assembled.name,
        "-o",
        str(output_file),
        "--from=markdown+raw_attribute",
        "--reference-doc=reference.docx",
        f"--resource-path={workspace}",
        "-V",
        "lang=es",
    ]
    subprocess.run(command, check=True, cwd=str(workspace))


def _is_figure_caption(para) -> bool:
    """Detect figure caption paragraphs (_Figura N. ..._): all runs italic, starts with 'Figura'."""
    text = para.text.strip()
    if not text.startswith("Figura "):
        return False
    text_runs = [r for r in para.runs if r.text.strip()]
    return bool(text_runs) and all(r.italic for r in text_runs)


def postprocess_doc(output_file: Path) -> None:
    doc = Document(output_file)

    # Apply table style and left-align all cell content.
    for table in doc.tables:
        try:
            table.style = "Table Grid"
        except Exception:
            try:
                table.style = "Table Grid Copilot"
            except Exception:
                pass
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    para.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Justify body paragraphs; center figure captions.
    for para in doc.paragraphs:
        if para.style.name.startswith("Heading"):
            continue
        if _is_figure_caption(para):
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    doc.save(output_file)


def build(reference_only: bool = False, keep_workspace: bool = False) -> int:
    ensure_reference_doc()
    if reference_only:
        print(f"Plantilla creada o verificada: {REFERENCE_DOC}")
        return 0

    if WORKSPACE.exists() and not keep_workspace:
        shutil.rmtree(WORKSPACE)

    includes = parse_includes()
    assembled, diagram_count = assemble_markdown(includes, WORKSPACE)
    run_pandoc(assembled, WORKSPACE, OUTPUT_FILE)
    postprocess_doc(OUTPUT_FILE)

    size_kib = OUTPUT_FILE.stat().st_size / 1024
    print(f"Documento generado: {OUTPUT_FILE}")
    print(f"Diagramas renderizados: {diagram_count}")
    print(f"Tamaño aproximado: {size_kib:.1f} KiB")
    return diagram_count


def main() -> int:
    parser = argparse.ArgumentParser(description="Genera el documento Word del panel.")
    parser.add_argument("--reference-only", action="store_true", help="Solo crea o verifica reference.docx.")
    parser.add_argument("--keep-workspace", action="store_true", help="Conserva la carpeta de trabajo sin espacios.")
    args = parser.parse_args()

    try:
        build(reference_only=args.reference_only, keep_workspace=args.keep_workspace)
    except subprocess.CalledProcessError as exc:
        print(f"Error ejecutando comando externo: {exc}", file=sys.stderr)
        return exc.returncode or 1
    except Exception as exc:
        print(f"Error: {exc}", file=sys.stderr)
        return 1

    return 0


if __name__ == "__main__":
    raise SystemExit(main())
